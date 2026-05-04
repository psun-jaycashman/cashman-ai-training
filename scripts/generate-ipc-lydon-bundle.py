"""
Generate the IPC Lydon Kansas Boiler project bundle used in
Module 4 Lesson 3 (Data Cleanup and Transformation).

Hypothetical project: install a 1.2 MMBtu/hr industrial boiler at the
"Sunflower Plains Ag Processing Facility" outside Topeka, Kansas. The
project is intentionally messy — different team members logged things in
different formats, the cost tracker was hand-edited multiple times, the
safety officer used a typewriter-style template, meeting transcripts came
out of three different apps, and the email thread is the usual chaos.

Files produced (under public/downloads/ipc-lydon-kansas-boiler/):
    cost-tracking.xlsx          PM-maintained: change orders, requisitions, cost tracking
    payment-log.xlsx            Accountant-maintained: pay requests + payments received
    superintendent-logs.docx    Daily field logs from two different supers, inconsistent format
    safety-officer-notes.docx   Daily safety notes (mostly), some missing days
    meeting-transcripts.txt     ~4 weekly team meetings, three different transcription styles
    project-emails.txt          Selected emails between PM, owner, sub, and accounting
    README.txt                  Index of files and the deliverables the user owes

A zip of the whole folder is written to:
    public/downloads/ipc-lydon-kansas-boiler.zip

Run:
    python3 scripts/generate-ipc-lydon-bundle.py
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, Inches, RGBColor


# --------------------------------------------------------------------------- #
# Project metadata (the "truth" these messy files allude to)
# --------------------------------------------------------------------------- #

PROJECT = {
    "name": "Sunflower Plains Boiler Install",
    "client": "Sunflower Plains Ag Processing, LLC",
    "location": "Topeka, KS (rural site, ~12 miles south)",
    "ipc_lydon_office": "IPC Lydon — Industrial Process Division (Stoughton, MA)",
    "scope": "Install one (1) 1.2 MMBtu/hr natural-gas package boiler with feedwater system, blowdown tank, stack, and integration to existing 150 psi steam header. Includes electrical (480V/3ph), controls (DCS tie-in), and 60-day commissioning support.",
    "contract_value": 8_400_000,
    "contract_signed": "2025-09-15",
    "ntp": "2025-10-06",
    "planned_substantial_completion": "2026-04-30",
    "actual_substantial_completion": "2026-06-12",
    "final_cost_at_writing": 9_865_000,
}

PROJECT_NUM = "IPC-2025-184"

PEOPLE = {
    "pm":          ("Karen Hayes",     "Project Manager (IPC Lydon)"),
    "super_a":     ("Dale Brennan",    "Field Superintendent A"),
    "super_b":     ("Mateo Ortiz",     "Field Superintendent B (relieved Dale wk 16)"),
    "safety":      ("Janelle Carter",  "Safety Officer"),
    "controls":    ("Wei Zhang",       "Controls / I&E Lead"),
    "accounting":  ("Tom Nichols",     "Accountant"),
    "owner_pm":    ("Brad Whitfield",  "Owner Rep — Sunflower Plains"),
    "owner_eng":   ("Linda Park",      "Owner Engineer"),
    "sub_mech":    ("Hawkeye Mechanical", "Subcontractor — pipefitting"),
    "sub_elec":    ("Plains Electric",  "Subcontractor — electrical"),
}


# --------------------------------------------------------------------------- #
# Styling helpers
# --------------------------------------------------------------------------- #

CASHMAN_GREEN = "2B6E2B"
LIGHT = "EFEFEF"
GRAY = "B7B7B7"

THIN = Side(style="thin", color=GRAY)
BOX = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
HEAD_FONT = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
GREEN_FILL = PatternFill("solid", fgColor=CASHMAN_GREEN)
ALT_FILL = PatternFill("solid", fgColor="F7F7F2")


# =========================================================================== #
# 1. cost-tracking.xlsx — PM-maintained cost tracker
# =========================================================================== #

CHANGE_ORDERS = [
    # CO#, Date, Description (reads like a working PM wrote it),
    # Status, Approved Amount, Notes
    ("CO-001", "2025-11-04", "Owner-requested upgrade: package boiler from 1.0 to 1.2 MMBtu/hr to handle planned dryer #3 expansion", "Approved",  240_000, "Per Brad email 11/3"),
    ("CO-002", "2025-11-21", "Differing site conditions — 4 ft of unmarked rebar in slab where boiler pad goes",                       "Approved",   78_500, "Field-verified by Dale"),
    ("CO-003", "2025-12-09", "Add insulation upgrade (from 1.5\" to 2\" calcium silicate) on steam header per owner engineer review",  "Approved",   54_000, "Linda Park request"),
    ("CO-004", "2026-01-14", "Re-route gas line — owner discovered abandoned ag chemical tank under planned route",                     "Approved",  186_000, "Required ESA review; 3-wk delay"),
    ("CO-005", "2026-01-28", "Owner request: tie-in DCS to MES system (was not in original scope)",                                   "Approved",  142_000, "Wei flagged as out of scope"),
    ("CO-006", "2026-02-18", "Replace blowdown tank — supplier shipped wrong material grade (304SS instead of 316SS)",                "Pending",    96_000, "Dispute with supplier; may be backcharge"),
    ("CO-007", "2026-03-09", "Crew overtime — owner requested accelerated schedule to recover from CO-004 delay",                     "Approved",  220_000, "Two-shift coverage 4 weeks"),
    ("CO-008", "2026-03-24", "Additional control valves on feedwater train — owner spec changed mid-fab",                             "Approved",   88_000, "Affected fab schedule"),
    ("CO-009", "2026-04-22", "Crane re-mob — high winds canceled first attempt, lost a full day plus rental",                         "Approved",   34_000, "Weather"),
    ("CO-010", "2026-05-12", "Punch list scope creep: owner added 17 items not in original closeout spec",                            "Pending",    62_000, "Negotiating"),
]

REQUISITIONS = [
    # Req#, Period, Submitted, Approved, Amount Requested, Approved $, Status
    ("REQ-1",  "Oct 2025",  "2025-11-05", "2025-11-19",   840_000,   840_000, "Paid"),
    ("REQ-2",  "Nov 2025",  "2025-12-04", "2025-12-22",  1_120_000, 1_080_000, "Paid (40k held — disputed mob)"),
    ("REQ-3",  "Dec 2025",  "2026-01-06", "2026-01-23",  1_240_000, 1_240_000, "Paid"),
    ("REQ-4",  "Jan 2026",  "2026-02-05", "2026-02-26",    980_000,   910_000, "Paid (70k disputed — CO-004 delay credit)"),
    ("REQ-5",  "Feb 2026",  "2026-03-04", "2026-03-24",  1_440_000, 1_440_000, "Paid"),
    ("REQ-6",  "Mar 2026",  "2026-04-08", "2026-04-29",  1_680_000, 1_680_000, "Paid"),
    ("REQ-7",  "Apr 2026",  "2026-05-06", "2026-05-26",  1_320_000, 1_280_000, "Paid (40k held — punch list)"),
    ("REQ-8",  "May 2026",  "2026-06-09", "Pending",       890_000,         0, "In review — disputed CO-006 + CO-010"),
]

COST_BUCKETS = [
    # Bucket, Original Budget, Forecast at Closeout, Variance
    ("Direct Labor",                 2_400_000, 2_980_000),
    ("Major Equipment (Boiler skid)", 1_900_000, 1_920_000),
    ("Subcontractors — Mechanical",  1_650_000, 1_910_000),
    ("Subcontractors — Electrical",    780_000,   840_000),
    ("Materials — Piping/Fittings",    520_000,   615_000),
    ("Materials — Insulation",         180_000,   235_000),
    ("Crane / Rigging",                240_000,   305_000),
    ("Permits / Fees",                  60_000,    72_000),
    ("PM / Engineering / Closeout",    370_000,   430_000),
    ("Contingency",                    300_000,   558_000),
]


def build_cost_tracking(out: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Change Orders"

    for col, w in {"A": 11, "B": 12, "C": 70, "D": 12, "E": 16, "F": 36}.items():
        ws.column_dimensions[col].width = w

    ws["A1"] = f"{PROJECT['name']} — PM Cost Tracker"
    ws["A1"].font = Font(size=14, bold=True)
    ws.merge_cells("A1:F1")
    ws["A2"] = f"Project: {PROJECT_NUM}  ·  Client: {PROJECT['client']}  ·  Maintained by: {PEOPLE['pm'][0]}"
    ws["A2"].font = Font(size=10, italic=True)
    ws.merge_cells("A2:F2")

    headers = ["CO #", "Date", "Description", "Status", "Approved $", "Notes"]
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=c, value=h)
        cell.font = HEAD_FONT
        cell.fill = GREEN_FILL
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = BOX

    for offset, row in enumerate(CHANGE_ORDERS):
        r = 5 + offset
        for c, v in enumerate(row, start=1):
            cell = ws.cell(row=r, column=c, value=v)
            cell.border = BOX
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if c == 5:
                cell.number_format = "$#,##0"
                cell.alignment = Alignment(horizontal="right", vertical="top")
            if r % 2 == 0:
                cell.fill = ALT_FILL

    # Requisitions sheet
    ws2 = wb.create_sheet("Requisitions")
    for col, w in {"A": 10, "B": 14, "C": 14, "D": 14, "E": 18, "F": 18, "G": 50}.items():
        ws2.column_dimensions[col].width = w
    ws2["A1"] = "Requisitions Submitted"
    ws2["A1"].font = Font(size=14, bold=True)
    ws2.merge_cells("A1:G1")
    headers2 = ["Req #", "Period", "Submitted", "Approved", "Requested $", "Approved $", "Status / Notes"]
    for c, h in enumerate(headers2, start=1):
        cell = ws2.cell(row=3, column=c, value=h)
        cell.font = HEAD_FONT
        cell.fill = GREEN_FILL
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = BOX
    for offset, row in enumerate(REQUISITIONS):
        r = 4 + offset
        for c, v in enumerate(row, start=1):
            cell = ws2.cell(row=r, column=c, value=v)
            cell.border = BOX
            if c in (5, 6):
                cell.number_format = "$#,##0"

    # Cost summary sheet — note that these often disagree slightly with reqs
    # (intentional: messy reality)
    ws3 = wb.create_sheet("Cost Summary")
    for col, w in {"A": 36, "B": 16, "C": 18, "D": 14}.items():
        ws3.column_dimensions[col].width = w
    ws3["A1"] = "Cost Bucket Summary (PM rolling forecast)"
    ws3["A1"].font = Font(size=14, bold=True)
    ws3.merge_cells("A1:D1")
    ws3["A2"] = "Note: this was hand-edited by the PM each month; see emails for caveats. Numbers do not exactly tie to requisitions."
    ws3["A2"].font = Font(size=9, italic=True, color="FF888888")
    ws3.merge_cells("A2:D2")
    headers3 = ["Cost Bucket", "Original Budget", "Forecast at Closeout", "Variance"]
    for c, h in enumerate(headers3, start=1):
        cell = ws3.cell(row=4, column=c, value=h)
        cell.font = HEAD_FONT
        cell.fill = GREEN_FILL
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = BOX
    total_orig = total_fc = 0
    for offset, (bucket, orig, fc) in enumerate(COST_BUCKETS):
        r = 5 + offset
        ws3.cell(row=r, column=1, value=bucket).border = BOX
        ws3.cell(row=r, column=2, value=orig).number_format = "$#,##0"
        ws3.cell(row=r, column=3, value=fc).number_format = "$#,##0"
        ws3.cell(row=r, column=4, value=fc - orig).number_format = "$#,##0"
        for c in range(1, 5):
            ws3.cell(row=r, column=c).border = BOX
        total_orig += orig
        total_fc += fc
    rT = 5 + len(COST_BUCKETS)
    ws3.cell(row=rT, column=1, value="TOTAL").font = Font(bold=True)
    ws3.cell(row=rT, column=2, value=total_orig).number_format = "$#,##0"
    ws3.cell(row=rT, column=3, value=total_fc).number_format = "$#,##0"
    ws3.cell(row=rT, column=4, value=total_fc - total_orig).number_format = "$#,##0"
    for c in range(1, 5):
        ws3.cell(row=rT, column=c).font = Font(bold=True)
        ws3.cell(row=rT, column=c).border = BOX
        ws3.cell(row=rT, column=c).fill = PatternFill("solid", fgColor=LIGHT)

    wb.save(out)


# =========================================================================== #
# 2. payment-log.xlsx — accountant's record
# =========================================================================== #

PAY_REQUESTS = [
    # Req#, Submit Date, Owner Approval Date, Amount Requested, Amount Received, Days to Pay, Notes
    ("REQ-1", "2025-11-05", "2025-11-26",   840_000,   840_000, 21, ""),
    ("REQ-2", "2025-12-04", "2026-01-08",  1_120_000, 1_080_000, 35, "40k withheld — owner disputed mob charge; resolved 1/8"),
    ("REQ-3", "2026-01-06", "2026-01-30",  1_240_000, 1_240_000, 24, ""),
    ("REQ-4", "2026-02-05", "2026-03-12",    980_000,   910_000, 35, "70k held against CO-004 delay credit"),
    ("REQ-5", "2026-03-04", "2026-04-01",  1_440_000, 1_440_000, 28, ""),
    ("REQ-6", "2026-04-08", "2026-05-07",  1_680_000, 1_680_000, 29, ""),
    ("REQ-7", "2026-05-06", "2026-06-15",  1_320_000, 1_280_000, 40, "40k held — punch list dispute (CO-010)"),
    ("REQ-8", "2026-06-09", "OPEN",           890_000,         0,  0, "Still in review — owner contesting CO-006 and CO-010"),
]


def build_payment_log(out: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Payment Log"
    for col, w in {"A": 10, "B": 14, "C": 18, "D": 16, "E": 16, "F": 12, "G": 60}.items():
        ws.column_dimensions[col].width = w

    ws["A1"] = f"{PROJECT['name']} — Payment Log"
    ws["A1"].font = Font(size=14, bold=True)
    ws.merge_cells("A1:G1")
    ws["A2"] = f"Maintained by: {PEOPLE['accounting'][0]} (Accounting)  ·  Project {PROJECT_NUM}"
    ws["A2"].font = Font(size=10, italic=True)
    ws.merge_cells("A2:G2")

    headers = ["Req #", "Submitted", "Owner Approved", "Requested $", "Received $", "Days to Pay", "Notes"]
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=c, value=h)
        cell.font = HEAD_FONT
        cell.fill = GREEN_FILL
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = BOX
    for offset, row in enumerate(PAY_REQUESTS):
        r = 5 + offset
        for c, v in enumerate(row, start=1):
            cell = ws.cell(row=r, column=c, value=v)
            cell.border = BOX
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if c in (4, 5):
                cell.number_format = "$#,##0"
                cell.alignment = Alignment(horizontal="right", vertical="top")
            if r % 2 == 0:
                cell.fill = ALT_FILL

    # Cash position summary
    rT = 5 + len(PAY_REQUESTS) + 1
    requested = sum(r[3] for r in PAY_REQUESTS)
    received = sum(r[4] for r in PAY_REQUESTS)
    ws.cell(row=rT, column=3, value="TOTAL").font = Font(bold=True)
    ws.cell(row=rT, column=4, value=requested).number_format = "$#,##0"
    ws.cell(row=rT, column=5, value=received).number_format = "$#,##0"
    ws.cell(row=rT + 1, column=3, value="OUTSTANDING").font = Font(bold=True, color="C00000")
    ws.cell(row=rT + 1, column=5, value=requested - received).number_format = "$#,##0"
    ws.cell(row=rT + 1, column=5).font = Font(bold=True, color="C00000")

    wb.save(out)


# =========================================================================== #
# 3. superintendent-logs.docx — daily field logs (intentionally messy)
# =========================================================================== #

# Two formats: Dale (terse, abbreviated, no consistent template) and Mateo
# (more structured but still freeform). Some days missing.

DALE_LOGS = [
    ("Mon 10/6/25",
     "NTP. Site walk w/ owner. Marked boiler pad location. Existing slab is gnarly — lots of patching. "
     "Hawkeye crew showed up but no fitter foreman, only 2 helpers. Owner says they want us moving by 10/13. "
     "Weather: 58F clear. Crew: 4. Notes: need to call Karen re: req schedule."),
    ("Wed 10/8/25",
     "Survey crew here. Set out boiler pad. CONFLICT — proposed pad location overlaps with abandoned cond drain line. "
     "Talked to Brad, he said \"just go around it.\" That's not how this works. Photos taken. Will write up in CO log. "
     "Crew: 4. Weather: rain pm."),
    ("Mon 10/13/25",
     "Excavation started for boiler pad. Hit something hard at 18\" depth. Stopped. Brought in Brad and his maintenance "
     "guy. Turns out it's a section of #8 rebar from an old equipment pad they pulled in like 2003. Not on any drawing. "
     "PHOTO of the rebar. Going to need to extract about 4ft of it. CO coming. Lost 1/2 day. Crew: 6."),
    ("Fri 10/24/25",
     "Pad pour today. 23 cy. 4500psi mix, 6\" thick. Pour started 0700 finished 1115. Cylinders pulled (3+3 pairs, 28 day). "
     "Weather perfect. Brad came by, was happy. Janelle did toolbox talk on hot weather work though it was only 72. Crew: 8."),
    ("Tue 11/4/25",
     "Got CO-001 from Karen — owner is upgrading boiler size. Means new submittal for the skid. Hawkeye foreman finally "
     "showed up. Talked through pipe routing for new size. Crew: 6."),
    ("Thu 11/20/25",
     "Owner Eng. (Linda P) here for site visit. Reviewed steam header insulation spec. She wants 2\" calcium silicate "
     "instead of 1.5\". Will affect schedule by ~1wk for material. CO needed. Karen aware. Weather: 28F!! Crew: 5."),
    ("Mon 12/8/25",
     "Pulled cond drain line out where pad goes. Took all morning, used the mini ex. Backfilled w/ flowable fill. "
     "Pad will get poured 12/10 weather permitting. Hawkeye fitters now 3 weeks behind on mobilizing big crew. "
     "Crew: 7."),
    ("Wed 1/7/26",
     "DISCOVERY. Surveying for gas line tie-in. Owner mtnce guy mentions \"oh yeah there used to be an ag chem tank "
     "underground around here.\" GREAT. Stopped all gas line work. Karen on phone w/ Brad. May need ESA. "
     "Crew: 6 (most stood down). Weather: 12F windy."),
    ("Mon 1/12/26",
     "Owner confirmed abandoned chem tank. ESA required before any further trenching. Will be 3wk delay min. "
     "Pulled gas crew off site. Re-tasking to controls panel rough-in. Crew: 4."),
    ("Wed 2/4/26",
     "ESA back, area is clean. Permit to re-trench in hand. Restarted gas line — different route, longer. CO-004 in. "
     "Crew: 6."),
    ("Fri 2/13/26",
     "Today is my last day on this job. Family thing back in MA. Mateo coming out to relieve me Mon. "
     "Walked him through site Thursday. Status: pad poured, boiler skid arrived (in storage), gas line 30%, "
     "controls panel 70%, electrical 50%, insulation not yet started, blowdown tank wrong material (returning). "
     "Crew: 8."),
]

MATEO_LOGS = [
    ("Monday, February 16, 2026",
     "Day 1 on site. Took over from Dale. Walked through with Brad (owner) and Wei (controls). "
     "Outstanding items: gas line completion, blowdown tank replacement, insulation install, controls programming, "
     "boiler set + alignment, commissioning. Estimated ~6 weeks remaining if no surprises.\n"
     "Crew today: 6 (3 IPC, 2 Hawkeye, 1 Plains Electric).\n"
     "Weather: 22F, light snow."),
    ("Wednesday, February 18, 2026",
     "Boiler skid on site at warehouse. Replacement blowdown tank ordered (correct 316SS this time). 3-week lead. "
     "Wei programming DCS, will tie into MES per CO-005.\n"
     "Crew: 7."),
    ("Wednesday, February 25, 2026",
     "Owner asking when boiler will set. I said 'after the new blowdown tank arrives (~3wk) and gas line is complete (~2wk).' "
     "Brad pushed for accelerated schedule. Karen working on pricing for OT and shift coverage.\n"
     "Crew: 8."),
    ("Friday, March 6, 2026",
     "Karen got CO-007 approved for accelerated schedule. Going to two shifts starting Monday 3/9. "
     "Need to coordinate with Hawkeye on second-shift pipefitters. May not have enough.\n"
     "Crew: 9."),
    ("Monday, March 23, 2026",
     "Boiler set today. Crane in at 0700. Set complete by 1100. Smooth. "
     "Photos to Karen. Owner happy. Janelle did pre-lift safety briefing.\n"
     "Crew: 12."),
    ("Tuesday, April 7, 2026",
     "Insulation install underway. Hawkeye made progress on piping over the weekend. "
     "Wei found an issue with the steam pressure transmitter signal — wrong scaling. Reflashing the DCS today.\n"
     "Crew: 11."),
    ("Friday, April 17, 2026",
     "First steam to header. Slow ramp per commissioning plan. Held for 4 hrs at 50psi, no leaks. "
     "Will hold overnight. Owner watching. Linda Park on site.\n"
     "Crew: 9."),
    ("Wednesday, April 29, 2026",
     "Substantial completion target was tomorrow (4/30) but we're ~2 wks behind due to the pressure transmitter rework "
     "and insulation crew shortage. Notified Karen and Brad. Punch list walk scheduled for next week.\n"
     "Crew: 8."),
    ("Thursday, May 14, 2026",
     "Punch walk with owner. They added 17 items not in original closeout spec — paint touch-ups in adjacent rooms, "
     "labels on existing equipment, etc. Karen will write up CO-010 for negotiation. Refused to sign substantial "
     "completion until resolved.\n"
     "Crew: 6."),
    ("Friday, June 12, 2026",
     "Substantial completion executed today. Punch list ~80% closed. Final 20% under negotiation per CO-010. "
     "Demob to start 6/15.\n"
     "Crew: 4."),
]


def build_super_logs(out: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading("Field Superintendent Daily Logs", level=0)
    h.runs[0].font.color.rgb = RGBColor(0x2B, 0x6E, 0x2B)

    p = doc.add_paragraph()
    p.add_run(f"Project: {PROJECT['name']} ({PROJECT_NUM})\n").bold = True
    p.add_run(f"Site: {PROJECT['location']}\n").italic = True
    p.add_run(f"Logs from {PEOPLE['super_a'][0]} (Oct 2025 – Feb 2026) and {PEOPLE['super_b'][0]} (Feb 2026 – Jun 2026)")

    doc.add_paragraph()
    doc.add_heading(f"Logs by {PEOPLE['super_a'][0]}", level=1)
    doc.add_paragraph(
        "(Dale wrote these on his phone in a notes app. Format is loose — no template, abbreviations everywhere.)"
    ).runs[0].italic = True

    for date_str, body in DALE_LOGS:
        p = doc.add_paragraph()
        p.add_run(date_str).bold = True
        doc.add_paragraph(body)

    doc.add_paragraph()
    doc.add_heading(f"Logs by {PEOPLE['super_b'][0]}", level=1)
    doc.add_paragraph(
        "(Mateo took over mid-February. He uses a more structured format but still writes in prose, not a spreadsheet template.)"
    ).runs[0].italic = True

    for date_str, body in MATEO_LOGS:
        p = doc.add_paragraph()
        p.add_run(date_str).bold = True
        doc.add_paragraph(body)

    doc.save(out)


# =========================================================================== #
# 4. safety-officer-notes.docx — daily safety notes (some missing days)
# =========================================================================== #

SAFETY_NOTES = [
    ("2025-10-06",  "Site orientation for IPC crew (4) and Hawkeye (3). Reviewed PPE, hot work, chemical exposure. No incidents."),
    ("2025-10-13",  "Toolbox: excavation safety. Discussed unmarked utilities, locate tickets. INCIDENT: laborer minor strain lifting jackhammer alone — used buddy lift after. First aid only, no medical."),
    ("2025-10-24",  "Pre-pour briefing. Hot weather plan (water, shade) even though only 72F. Concrete chemical splash protocol reviewed. No incidents."),
    ("2025-11-04",  "Toolbox: confined space refresher in advance of pipe work. Documented attendees: 6."),
    ("2025-11-13",  "Near-miss: scaffold board flipped under foreman, no fall, no injury. Tagged scaffold for inspection. Scaffold sub on site by EOD; replaced 4 boards."),
    ("2025-12-02",  "Cold weather toolbox. Below-30 work plan reviewed. Heated break trailer arrived on site."),
    # Missing: 2025-12-15 (Janelle out — owner's safety guy covered, no notes left)
    ("2026-01-07",  "STAND-DOWN. Discovery of unmarked abandoned ag chem tank during gas line trenching. All personnel pulled from area pending ESA. Tank vented, capped years ago per owner. No exposure. Briefed crew on procedure."),
    ("2026-01-12",  "Continued stand-down. Reviewed personal exposure protocols with crew while off the gas line scope. Used time for hot work refresher."),
    ("2026-02-04",  "Trenching restart. Reviewed shoring requirements (>5 ft trench), confined space entry protocol. No incidents."),
    # Missing: 2026-02-13 (Janelle was at quarterly safety conference)
    ("2026-02-25",  "Welcomed Mateo (new super) to safety program. Walked sitewide hazards: open trench, energized panel, overhead crane path. No incidents."),
    ("2026-03-09",  "Two-shift work begins. Reviewed fatigue management, shift handover protocols. New 2nd-shift crew (3) oriented."),
    ("2026-03-23",  "BIG LIFT TODAY: boiler set. Pre-lift briefing 0630. Critical lift plan reviewed and signed. Site cordoned. Crane operator certs verified. Lift completed 1100, no incidents. PHOTO of pre-lift huddle."),
    ("2026-04-07",  "INCIDENT: insulation worker contacted hot pipe stub during demo. Minor 1st-degree burn to forearm, 4 sq in. First aid on site. Refused medical. Reported to OSHA log per protocol. Root cause: insulator did not lockout per LOTO procedure for a stub he assumed was cold. Stand-down for crew, full LOTO refresher 4/8."),
    ("2026-04-08",  "Mandatory LOTO refresher for all hot work. 11 attendees. Updated procedure to require physical tag verification by 2nd person."),
    ("2026-04-17",  "First steam to header. Reviewed steam burn risk, exclusion zones, communication protocols (radio check before any valve operation). No incidents."),
    ("2026-04-29",  "Substantial completion walk. Reviewed open punch items for safety implications. None require stop-work. Owner punch added items in adjacent occupied space — coordinated access protocols with Brad."),
    ("2026-05-14",  "Punch walk supplemental. Adjacent-space work requires coordination with owner ops. Reviewed hot work permit for any cutting/welding."),
    ("2026-06-12",  "Demob safety briefing. Crane removal scheduled 6/15. No incidents during punch list closeout phase."),
]


def build_safety_notes(out: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"  # different look from super logs
    style.font.size = Pt(11)

    h = doc.add_heading("Safety Officer Daily Notes", level=0)
    h.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p = doc.add_paragraph()
    p.add_run(f"Officer: {PEOPLE['safety'][0]}\n").bold = True
    p.add_run(f"Project: {PROJECT['name']} ({PROJECT_NUM})\n")
    p.add_run("Notes are entered same-day except where Janelle was offsite (a few entries missing).").italic = True

    doc.add_paragraph()

    # Use a simple table layout because Janelle is methodical
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Date"
    hdr[1].text = "Notes"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for date, note in SAFETY_NOTES:
        row = table.add_row().cells
        row[0].text = date
        row[1].text = note

    doc.save(out)


# =========================================================================== #
# 5. meeting-transcripts.txt — three different transcript styles
# =========================================================================== #

MEETING_TRANSCRIPTS = """\
================================================================================
CASHMAN / IPC LYDON — SUNFLOWER PLAINS BOILER (IPC-2025-184)
WEEKLY OWNER STATUS MEETINGS — PARTIAL TRANSCRIPTS
================================================================================
This file aggregates four weekly status meeting transcripts captured over the
life of the project. Three different transcription tools were used (Otter, Zoom
auto-transcript, and one set of hand-typed notes), so formats vary.


--------------------------------------------------------------------------------
MEETING 1 — Tuesday, October 14, 2025  (Otter export)
Attendees (per Otter): Karen Hayes (IPC), Dale Brennan (IPC), Brad Whitfield (Owner),
Linda Park (Owner Eng), Tom Nichols (Acct, dial-in)
--------------------------------------------------------------------------------

00:00:12  Karen:  thanks everybody. quick walk through where we are. dale you want to
                  start with field?
00:00:24  Dale:   yeah so um pad excavation hit the rebar on monday like i mentioned.
                  brad you saw the photo right
00:00:34  Brad:   i did. i didn't know that was there
00:00:38  Dale:   yeah it's not on any drawing your maintenance guy stan
                  said it was from the old equipment pad you all pulled in 03 i think
00:00:51  Brad:   stan would know
00:00:54  Karen:  so we'll write that up as a CO. probably 70-ish thousand
00:01:02  Brad:   approved subject to seeing the breakdown
00:01:10  Linda:  on insulation i've been thinking we should bump to 2 inch on the steam
                  header. the existing piping in that gallery runs hot
00:01:22  Karen:  okay we can price that. probably another CO
00:01:28  Dale:   that's gonna affect schedule by like a week for material
00:01:33  Brad:   acceptable
00:01:38  Karen:  tom anything from accounting
00:01:41  Tom:    [audio clipped] req 1 went out friday should hit 21 days
00:01:48  Karen:  okay any other items? no? thanks all


--------------------------------------------------------------------------------
MEETING 2 — Tuesday, January 13, 2026  (Zoom auto-transcript, raw)
--------------------------------------------------------------------------------

[Karen Hayes]: All right, I think we have everybody. Let's start. Dale,
do you want to give us the field update?

[Dale Brennan]: Yeah uh so as everybody knows on Wednesday we hit the
chem tank. Or, the place where there used to be a chem tank. The owner
maintenance team says it was vented and capped maybe 15 years ago. We
stood the gas line crew down. We've got an ESA scheduled for Thursday.

[Brad Whitfield]: Karen, what's the schedule impact on this.

[Karen Hayes]: Best case three weeks. We need the ESA, then if it's
clean we restart but on a different route. If it's not clean we're
talking permits and remediation. I want to be honest, that could be
two months.

[Brad Whitfield]: Okay. I need a CO write up for the re-route. And we
need to start talking about whether we hold to the April substantial.

[Karen Hayes]: I think we need to assume we slip. I don't want to
commit to a new date until we know about the contamination.

[Wei Zhang]: Just a heads up, while the gas line is on hold I'm going
to push controls panel rough in. We can keep moving.

[Brad Whitfield]: Good. Linda anything from your end.

[Linda Park]: I want to revisit the DCS to MES tie in. We talked about
it being out of scope but I think we need it now. Can we get a price.

[Karen Hayes]: Wei can scope it, I'll have a CO by Friday.

[Brad Whitfield]: Approved subject to price.

[Karen Hayes]: Okay, Tom?

[Tom Nichols]: Req 3 paid on time. Req 2 was held up because they
disputed the mob, that's been resolved. Heads up on req 4 which will
include the CO 002 reconciliation.

[Brad Whitfield]: Any safety items?

[Janelle Carter]: One. Standdown was clean, no exposure. Crew using the
downtime for hot work refresher.

[Brad Whitfield]: Good. Same time next week.


--------------------------------------------------------------------------------
MEETING 3 — Tuesday, March 24, 2026  (Hand-typed notes by Karen)
--------------------------------------------------------------------------------

WEEKLY MEETING NOTES                                   3/24/26
Attending: K. Hayes, M. Ortiz, B. Whitfield, W. Zhang, J. Carter, T. Nichols (phone),
           L. Park (left at 3:15)

1. BOILER SET — completed yesterday morning. Lift went smoothly, 0700-1100.
   Mateo handed photos to Brad.

2. OUTSTANDING SCOPE
   - Insulation: ~50% complete, Hawkeye crew shortage continues. Mateo
     working with Hawkeye PM to add bodies for 2 weeks.
   - Controls/DCS: 80%. Wei waiting on owner IT for MES tie-in credentials.
     Brad will push internally.
   - Blowdown tank: replacement 316SS arrived 3/22, install this week.
   - Punch list: not yet started, walk planned for first week of May.

3. SCHEDULE
   - Substantial completion target: 4/30 (was, now 5/8 estimated)
   - Mateo: this is conditional on no further surprises and Wei getting MES creds.
   - Brad: would like to know by 4/15 if 5/8 is realistic.

4. CHANGE ORDERS
   - CO-007 (accelerated schedule OT) approved 3/9. Two-shift coverage active.
   - CO-008 (additional control valves) approved 3/24 (today).
   - CO-006 (blowdown tank wrong material) STILL PENDING. Karen escalating
     to supplier. Owner does not want to absorb this.

5. ACCOUNTING (Tom)
   - Req 5 went out 3/4, paid 4/1, 28 days.
   - Req 6 will be larger due to OT and CO-007 first reconciliation.

6. SAFETY (Janelle)
   - Two-shift coverage: no incidents to date. Continuing fatigue management
     monitoring.

7. ACTION ITEMS
   - Brad: push owner IT for MES creds
   - Karen: escalate CO-006 supplier dispute
   - Mateo: coordinate insulation crew add
   - Wei: submit CO-008 detail
   - Tom: send Req 6 forecast to Brad


--------------------------------------------------------------------------------
MEETING 4 — Tuesday, May 19, 2026  (Otter export)
--------------------------------------------------------------------------------

00:00:08  Karen:  okay let's go. punch list status mateo
00:00:14  Mateo:  we've closed about 70 percent. the original closeout list. on the
                  17 items the owner added we've completed 4. 13 still open
00:00:28  Brad:   why are we still talking about the 17. those weren't in the contract
00:00:33  Mateo:  that's the CO-010 issue
00:00:36  Karen:  brad we sent the CO-010 detail two weeks ago. we have to negotiate
                  what's in scope before we can keep working on items that aren't
00:00:48  Brad:   look, we need this thing closed out. some of those items affect our
                  ability to start the dryer 3 expansion
00:00:57  Karen:  i hear you. let me work with you offline on a compromise. but we
                  cant absorb 60k of unbudgeted punch
00:01:08  Brad:   understood. let's talk after this
00:01:13  Karen:  blowdown tank CO-006 also still pending on supplier. our position is
                  the supplier eats it. theirs is that the spec was ambiguous. lawyers
                  involved now. dont expect resolution this month
00:01:31  Tom:    req 7 came in 40k short on the punch list withhold
00:01:36  Karen:  yes that ties to CO-010 negotiation
00:01:42  Janelle: no incidents. punch walk safe. demob plan being drafted
00:01:49  Karen:  okay anything else? no. talk next week
"""


def build_meeting_transcripts(out: Path) -> None:
    out.write_text(MEETING_TRANSCRIPTS, encoding="utf-8")


# =========================================================================== #
# 6. project-emails.txt — saved selected emails (RFC 822-ish style)
# =========================================================================== #

EMAILS = """\
================================================================================
SAVED EMAILS — IPC-2025-184 (SUNFLOWER PLAINS BOILER)
Exported from Outlook by Karen Hayes for project archive. Selected highlights only.
================================================================================


From: Brad Whitfield <bwhitfield@sunflowerplains.com>
To: Karen Hayes <khayes@ipclydon.com>
Date: 2025-11-03 16:42
Subject: Boiler size

Karen,

Talked to my CFO. We need the bigger 1.2 unit. The dryer 3 expansion is moving
forward and there's no point putting in a 1.0 if we'll outgrow it in 18 months.

What's the cost delta? Send a CO.

Brad

--

From: Karen Hayes <khayes@ipclydon.com>
To: Brad Whitfield <bwhitfield@sunflowerplains.com>
Cc: Tom Nichols <tnichols@ipclydon.com>
Date: 2025-11-04 09:11
Subject: RE: Boiler size

Brad,

Per our convo, here's the breakdown for upgrading from 1.0 to 1.2 MMBtu/hr:

  Boiler skid delta:           $185,000
  Larger feedwater piping:      $32,000
  Heavier-gauge stack:          $14,000
  Re-engineering / submittal:    $9,000
                            ---------
  Total:                       $240,000

CO-001 attached. Schedule impact ~10 days for resubmittal and re-fab.

Karen

--

From: Karen Hayes <khayes@ipclydon.com>
To: project-team@ipclydon.com
Date: 2026-01-07 17:23
Subject: STAND DOWN — Sunflower Plains gas line

Team,

Crew hit an unmarked abandoned chemical tank during gas line trenching today.
Owner's maintenance team says it was vented and capped ~15 years ago. No
exposure observed.

We have ALL gas line crew off site until ESA is complete. Targeting late
next week.

Will brief at Tuesday's status meeting. Expect a CO and a schedule
slip — magnitude TBD.

Karen

--

From: Linda Park <linda.park@sunflowerplains.com>
To: Wei Zhang <wzhang@ipclydon.com>
Cc: Karen Hayes <khayes@ipclydon.com>; Brad Whitfield <bwhitfield@sunflowerplains.com>
Date: 2026-01-15 11:07
Subject: DCS to MES integration

Wei,

Following up on Tuesday's meeting. We need the boiler DCS to push the
following tags to our MES every 5 seconds:

  - Steam pressure (header)
  - Steam temperature (header)
  - Feedwater flow
  - Stack O2 / CO
  - Fuel flow (gas)
  - Burner state
  - Alarms (any)

This was not in the original scope but it's important for our heat
balance reporting. Please scope and price.

Linda

--

From: Tom Nichols <tnichols@ipclydon.com>
To: Karen Hayes <khayes@ipclydon.com>
Date: 2026-02-26 14:55
Subject: Req 4 reconciliation — Sunflower

Karen,

Owner approved Req 4 today, but withheld $70k against CO-004 delay. Their
position is that the chem tank was on their property and they bear that
risk, but they're claiming we should have phase-loaded the gas crew
elsewhere to avoid full standby cost. Disputed; we have the daily logs
showing crew was redirected to controls panel work where possible.

Tom

--

From: Mateo Ortiz <mortiz@ipclydon.com>
To: Karen Hayes <khayes@ipclydon.com>
Date: 2026-04-07 18:14
Subject: Burn incident today — Sunflower

Karen,

Heads up. An insulator from Hawkeye contacted a hot pipe stub during demo
this afternoon. Minor 1st-degree burn forearm, ~4 sq in. First aid on
site, refused medical, refused trip to clinic. We've logged it on the
OSHA recordable.

Janelle is doing a stand-down tomorrow morning for full LOTO refresher.

Root cause: insulator did not verify isolation on what he assumed was a
cold stub. Procedure update: physical tag verification by a second
person before any hot work near boiler.

Mateo

--

From: Brad Whitfield <bwhitfield@sunflowerplains.com>
To: Karen Hayes <khayes@ipclydon.com>
Date: 2026-05-15 09:02
Subject: Punch list

Karen,

I know we are at odds on CO-010. The 17 items I added at walkthrough are
things our maintenance team needs in order to take care of this thing
once you all leave. Some are direct safety items like missing labels on
adjacent equipment. I am willing to compromise on the paint and the
non-safety items, but the safety-related ones I expect you to absorb.

Let's talk Friday.

Brad

--

From: Karen Hayes <khayes@ipclydon.com>
To: Brad Whitfield <bwhitfield@sunflowerplains.com>
Date: 2026-05-19 16:33
Subject: RE: Punch list

Brad,

Reviewing your 17 items, I count:
   - 4 items that are safety-related and adjacent to our scope.
     We will absorb these.
   - 6 items that are safety-related but in spaces we never worked
     in. These are not our scope.
   - 7 items that are aesthetic/labeling. Not our scope.

I propose: we close the 4 at our cost and submit a revised CO-010 for
just the 6 safety items in adjacent spaces, at half rate as a goodwill
gesture given the project length.

Total revised CO-010 ask: $34,000 (down from $62,000).

Let me know.

Karen
"""


def build_emails(out: Path) -> None:
    out.write_text(EMAILS, encoding="utf-8")


# =========================================================================== #
# 7. README.txt — index + deliverables
# =========================================================================== #

README = """\
IPC LYDON — SUNFLOWER PLAINS BOILER PROJECT (Project IPC-2025-184)
===================================================================
Hypothetical project bundle for the Module 4, Lesson 3 exercise
(Data Cleanup and Transformation).

The story
---------
IPC Lydon — the industrial process / mechanical contracting subsidiary —
was hired to install a new 1.2 MMBtu/hr natural-gas package boiler at the
Sunflower Plains Ag Processing Facility outside Topeka, Kansas. Contract
value $8.4M, contract signed September 2025, NTP October 2025, planned
substantial completion April 30, 2026.

It came in at ~$9.86M, substantial completion June 12, 2026. There were
ten change orders, several disputes, an OSHA-recordable burn, an
abandoned chemical tank discovery, a wrong-material blowdown tank, and
a contentious 17-item punch list addition by the owner.

The project involved many people and they all kept their own records in
their own formats:

Files included
--------------
1. cost-tracking.xlsx
   PM-maintained workbook with three sheets: Change Orders, Requisitions,
   Cost Bucket Summary. Numbers between sheets do not exactly tie — the
   PM was hand-editing.

2. payment-log.xlsx
   Accountant-maintained log of every requisition submitted and what was
   paid against it, with dispute notes.

3. superintendent-logs.docx
   Daily field logs from two superintendents:
     - Dale Brennan (Oct 2025 – mid-Feb 2026): terse, abbreviated, no template
     - Mateo Ortiz (mid-Feb 2026 – Jun 2026): more structured prose
   Some days are missing.

4. safety-officer-notes.docx
   Daily safety officer notes from Janelle Carter, in a Word table.
   A couple of days are missing (officer offsite).

5. meeting-transcripts.txt
   Four weekly owner status meetings. Three different transcription
   tools were used; format varies wildly.

6. project-emails.txt
   Selected emails between PM, owner, controls lead, accounting, and
   the new super.

What you owe
------------
Use ChatGPT or the Cashman AI Portal to digest these files. Then produce:

  REQUIRED:
  (1) An overall PROJECT SUMMARY AND LOOK-BACK in Word (one page).
      Cover the scope, what went well, what went wrong, root causes,
      and the final financial picture.

  (2) A SEQUENCE OF EVENTS table in Excel suitable for an audit.
      One row per significant event: date, event, source document
      (which file mentions it), $ impact if applicable, schedule
      impact in days, who decided.

  BONUS:
  (3) A LESSONS LEARNED PowerPoint (5-10 slides) you could give to
      another PM about to start a similar industrial project.

Submit your work into the shared library in busibox so other team
members can see what you produced. The lesson page has the link.

Tips for working with AI
------------------------
- Upload the files one by one (or all together if your tool supports it).
- Ask for a SINGLE-PAGE summary, not a regurgitation. AI will
  default to long. Ask it to compress.
- For the audit table, be explicit: "Build a table with these
  columns: Date, Event, Source File, $ Impact, Schedule Days, Decision Owner."
- Have AI cross-check the cost tracker against the payment log —
  these don't match exactly, and the discrepancies are the story.
- Ask AI WHY this project went over, not just what happened.
"""


def build_readme(out: Path) -> None:
    out.write_text(README, encoding="utf-8")


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #


def main() -> None:
    repo_root = Path(__file__).resolve().parent.parent
    bundle_dir = repo_root / "public" / "downloads" / "ipc-lydon-kansas-boiler"
    bundle_dir.mkdir(parents=True, exist_ok=True)

    builders = [
        ("cost-tracking.xlsx",        build_cost_tracking),
        ("payment-log.xlsx",          build_payment_log),
        ("superintendent-logs.docx",  build_super_logs),
        ("safety-officer-notes.docx", build_safety_notes),
        ("meeting-transcripts.txt",   build_meeting_transcripts),
        ("project-emails.txt",        build_emails),
        ("README.txt",                build_readme),
    ]

    for filename, fn in builders:
        out = bundle_dir / filename
        fn(out)
        print(f"  ✓ {out.relative_to(repo_root)}")

    # Bundle everything as a zip for one-click download.
    zip_path = repo_root / "public" / "downloads" / "ipc-lydon-kansas-boiler.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for filename, _ in builders:
            zf.write(bundle_dir / filename, arcname=f"ipc-lydon-kansas-boiler/{filename}")
    print(f"  ✓ {zip_path.relative_to(repo_root)}")
    print(f"\nProject {PROJECT_NUM}: contract ${PROJECT['contract_value']:,}, "
          f"final ${PROJECT['final_cost_at_writing']:,}")


if __name__ == "__main__":
    main()
